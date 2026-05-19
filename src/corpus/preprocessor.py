"""
Corpus preprocessor — handles document ingestion, text extraction,
hierarchical parsing, and metadata enrichment for all domains.

Supports PDF, DOCX, TXT, and Markdown formats.
Outputs a standardized JSON structure that all pipelines consume.
"""

from __future__ import annotations


import json
import os
import hashlib
from pathlib import Path
from dataclasses import dataclass, field, asdict
from typing import Any

from loguru import logger


@dataclass
class Section:
    """A single section within a document's hierarchy."""

    id: str                           # e.g., "3.2.1"
    title: str
    content: str
    level: int                        # 0=root, 1=chapter, 2=section, ...
    page_start: int | None = None
    page_end: int | None = None
    word_count: int = 0
    children: list["Section"] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)

    def __post_init__(self):
        self.word_count = len(self.content.split())


@dataclass
class ProcessedDocument:
    """A fully processed document ready for all pipelines."""

    doc_id: str
    filename: str
    domain: str                       # finance | legal | technical
    title: str
    full_text: str                    # Flat text (for BM25)
    sections: list[Section]           # Hierarchical structure (for tree RAG)
    total_pages: int | None = None
    total_words: int = 0
    metadata: dict = field(default_factory=dict)

    def __post_init__(self):
        self.total_words = len(self.full_text.split())
        if not self.doc_id:
            self.doc_id = hashlib.md5(self.filename.encode()).hexdigest()[:12]


@dataclass
class ProcessedCorpus:
    """Collection of processed documents for a single domain."""

    domain: str
    documents: list[ProcessedDocument]
    total_documents: int = 0
    total_words: int = 0

    def __post_init__(self):
        self.total_documents = len(self.documents)
        self.total_words = sum(d.total_words for d in self.documents)


class CorpusPreprocessor:
    """
    Multi-format document preprocessor that outputs a standardized
    hierarchical + flat text representation for all RAG pipelines.

    Usage:
        preprocessor = CorpusPreprocessor()
        corpus = preprocessor.process("data/raw/finance/", domain="finance")
        preprocessor.save(corpus, "data/processed/finance/")
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}

    def __init__(self, config: dict[str, Any] | None = None):
        self.config = config or {}

    def process(self, raw_path: str, domain: str) -> ProcessedCorpus:
        """
        Process all documents in a directory.

        Args:
            raw_path: Path to directory of raw documents.
            domain: Domain label (finance, legal, technical).

        Returns:
            ProcessedCorpus with all documents parsed.
        """
        raw_dir = Path(raw_path)
        if not raw_dir.exists():
            raise FileNotFoundError(f"Raw corpus path not found: {raw_path}")

        documents = []
        files = sorted(
            f for f in raw_dir.iterdir()
            if f.suffix.lower() in self.SUPPORTED_EXTENSIONS
        )

        logger.info(f"Processing {len(files)} documents from {raw_path} (domain={domain})")

        for filepath in files:
            try:
                doc = self._process_file(filepath, domain)
                documents.append(doc)
                logger.debug(f"  ✓ {filepath.name}: {doc.total_words} words, "
                             f"{len(doc.sections)} top-level sections")
            except Exception as e:
                logger.error(f"  ✗ Failed to process {filepath.name}: {e}")

        corpus = ProcessedCorpus(domain=domain, documents=documents)
        logger.info(
            f"Corpus ready: {corpus.total_documents} docs, "
            f"{corpus.total_words:,} total words"
        )
        return corpus

    def _process_file(self, filepath: Path, domain: str) -> ProcessedDocument:
        """Process a single file based on its extension."""
        ext = filepath.suffix.lower()

        if ext == ".pdf":
            return self._process_pdf(filepath, domain)
        elif ext in (".txt", ".md"):
            return self._process_text(filepath, domain)
        elif ext == ".docx":
            return self._process_docx(filepath, domain)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    def _process_pdf(self, filepath: Path, domain: str) -> ProcessedDocument:
        """Extract text and structure from PDF."""
        from pypdf import PdfReader

        reader = PdfReader(str(filepath))
        pages_text = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            pages_text.append(text)

        full_text = "\n\n".join(pages_text)
        sections = self._extract_sections_from_text(full_text, filepath.name)

        return ProcessedDocument(
            doc_id="",
            filename=filepath.name,
            domain=domain,
            title=filepath.stem.replace("_", " ").title(),
            full_text=full_text,
            sections=sections,
            total_pages=len(reader.pages),
            metadata={"source_format": "pdf", "source_path": str(filepath)},
        )

    def _process_text(self, filepath: Path, domain: str) -> ProcessedDocument:
        """Extract text and structure from plain text / markdown."""
        full_text = filepath.read_text(encoding="utf-8", errors="replace")
        sections = self._extract_sections_from_text(full_text, filepath.name)

        return ProcessedDocument(
            doc_id="",
            filename=filepath.name,
            domain=domain,
            title=filepath.stem.replace("_", " ").title(),
            full_text=full_text,
            sections=sections,
            metadata={"source_format": filepath.suffix, "source_path": str(filepath)},
        )

    def _process_docx(self, filepath: Path, domain: str) -> ProcessedDocument:
        """Extract text and structure from DOCX."""
        from docx import Document

        doc = Document(str(filepath))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)
        sections = self._extract_sections_from_text(full_text, filepath.name)

        return ProcessedDocument(
            doc_id="",
            filename=filepath.name,
            domain=domain,
            title=filepath.stem.replace("_", " ").title(),
            full_text=full_text,
            sections=sections,
            metadata={"source_format": "docx", "source_path": str(filepath)},
        )

    def _extract_sections_from_text(
        self, text: str, filename: str
    ) -> list[Section]:
        """
        Parse text into hierarchical sections using heading detection.

        Handles:
        - Markdown headings (# ## ###)
        - Numbered headings (1. 1.1 1.1.1)
        - ALL-CAPS headings (common in legal/financial docs)
        """
        lines = text.split("\n")
        sections = []
        current_section = None
        section_counter = [0]  # Mutable counter for ID generation

        content_buffer = []

        for line in lines:
            heading_level = self._detect_heading_level(line)

            if heading_level is not None and heading_level >= 1:
                # Save previous section
                if current_section is not None:
                    current_section.content = "\n".join(content_buffer).strip()
                    sections.append(current_section)

                section_counter[0] += 1
                current_section = Section(
                    id=str(section_counter[0]),
                    title=self._clean_heading(line),
                    content="",
                    level=heading_level,
                )
                content_buffer = []
            else:
                content_buffer.append(line)

        # Don't forget the last section
        if current_section is not None:
            current_section.content = "\n".join(content_buffer).strip()
            sections.append(current_section)
        elif content_buffer:
            # No headings found — treat entire doc as one section
            sections.append(Section(
                id="1",
                title=filename,
                content="\n".join(content_buffer).strip(),
                level=0,
            ))

        return sections

    def _detect_heading_level(self, line: str) -> int | None:
        """Detect if a line is a heading and return its level."""
        stripped = line.strip()
        if not stripped:
            return None

        # Markdown headings
        if stripped.startswith("#"):
            level = 0
            for ch in stripped:
                if ch == "#":
                    level += 1
                else:
                    break
            return min(level, 6)

        # ALL-CAPS headings (>3 words, all uppercase)
        words = stripped.split()
        if (len(words) >= 2 and len(words) <= 10
                and stripped == stripped.upper()
                and stripped[0].isalpha()):
            return 1

        return None

    def _clean_heading(self, line: str) -> str:
        """Clean a heading line for use as a section title."""
        return line.strip().lstrip("#").strip()

    def save(self, corpus: ProcessedCorpus, output_path: str) -> None:
        """
        Save processed corpus to disk as JSON.

        Args:
            corpus: ProcessedCorpus to serialize.
            output_path: Directory to write to.
        """
        out_dir = Path(output_path)
        out_dir.mkdir(parents=True, exist_ok=True)

        # Save each document individually
        for doc in corpus.documents:
            doc_path = out_dir / f"{doc.doc_id}.json"
            with open(doc_path, "w", encoding="utf-8") as f:
                json.dump(asdict(doc), f, indent=2, ensure_ascii=False)

        # Save corpus manifest
        manifest = {
            "domain": corpus.domain,
            "total_documents": corpus.total_documents,
            "total_words": corpus.total_words,
            "documents": [
                {"doc_id": d.doc_id, "filename": d.filename, "words": d.total_words}
                for d in corpus.documents
            ],
        }
        manifest_path = out_dir / "manifest.json"
        with open(manifest_path, "w", encoding="utf-8") as f:
            json.dump(manifest, f, indent=2)

        logger.info(f"Saved corpus to {output_path} ({corpus.total_documents} docs)")

    def load(self, processed_path: str) -> ProcessedCorpus:
        """Load a previously processed corpus from disk."""
        proc_dir = Path(processed_path)
        manifest_path = proc_dir / "manifest.json"

        with open(manifest_path, "r") as f:
            manifest = json.load(f)

        documents = []
        for doc_info in manifest["documents"]:
            doc_path = proc_dir / f"{doc_info['doc_id']}.json"
            with open(doc_path, "r") as f:
                data = json.load(f)
            # Reconstruct Section objects
            sections = [Section(**s) for s in data.get("sections", [])]
            data["sections"] = sections
            documents.append(ProcessedDocument(**data))

        return ProcessedCorpus(
            domain=manifest["domain"],
            documents=documents,
        )
